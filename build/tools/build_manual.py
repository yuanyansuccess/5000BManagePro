# -*- coding: utf-8 -*-
"""
生成《5000B 管理系统 环境搭建手册》保姆级 Word 文档。

用法（项目根目录执行）：
    python build/tools/build_manual.py

输出：build/5000B管理系统环境搭建手册.docx
"""
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = r"D:\5000\5000BManagePro"
OUT = os.path.join(ROOT, "build", "5000B管理系统环境搭建手册.docx")


def set_cjk(run, name="宋体"):
    """设置中文字体（eastAsia），避免中文显示为默认西文字体。"""
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), name)


def shade(paragraph, fill="F2F2F2"):
    """给段落加底纹（用于代码块）。"""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    paragraph._p.get_or_add_pPr().append(shd)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(26)
    set_cjk(r, "宋体")
    return p


def add_sub(doc, text, size=14, color=None, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = color
    set_cjk(r)
    return p


def add_h(doc, text, level=1):
    sizes = {1: 16, 2: 13, 3: 11.5}
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(sizes.get(level, 11.5))
    r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x73)
    set_cjk(r)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_p(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    set_cjk(r)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_steps(doc, items):
    """有序步骤列表。"""
    for i, t in enumerate(items, 1):
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(t)
        r.font.size = Pt(10.5)
        set_cjk(r)


def add_bullets(doc, items):
    for t in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(t)
        r.font.size = Pt(10.5)
        set_cjk(r)


def add_code(doc, text):
    for line in text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(9)
        shade(p)


def add_tip(doc, text, label="提示"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    r1 = p.add_run("【%s】" % label)
    r1.bold = True
    r1.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    set_cjk(r1)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    set_cjk(r2)
    return p


def add_table(doc, header, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        set_cjk(r)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(9)
            set_cjk(r)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    return t


def main():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(10.5)
    normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")

    # ============ 封面 ============
    add_title(doc, "5000B 管理系统")
    add_sub(doc, "环 境 搭 建 手 册", size=20, bold=True)
    add_sub(doc, "保姆级教程 —— 从一台空机器到系统跑起来", size=12,
            color=RGBColor(0x66, 0x66, 0x66))
    doc.add_paragraph()
    add_sub(doc, "版本：V1.0", size=11)
    add_sub(doc, "适用对象：零基础新手（无需开发经验，照着做即可）", size=11)
    add_sub(doc, "预计耗时：约 60 ~ 90 分钟（含软件下载）", size=11)
    add_sub(doc, "编制日期：2026-09-02", size=11)
    doc.add_page_break()

    # ============ 第0章 阅读说明 ============
    add_h(doc, "第 0 章  阅读说明", 1)
    add_p(doc, "本手册面向从未接触过本项目的新同事，目标是：只看这一份文档，就能把系统装好、跑起来、并生成出第一份《软件开发计划》文档。")
    add_steps(doc, [
        "严格按章节顺序操作，不要跳步；每一步都给出了可直接复制的命令。",
        "灰色底纹框里的内容是“需要你输入的命令”，复制粘贴即可，不要手动敲。",
        "【提示】是容易踩坑的地方，请务必看一眼。",
        "每完成一大章，可以运行一次自检命令，确认这一步是否成功。",
        "遇到报错先翻到第 11 章《常见问题排查》，80% 的问题都能在那里找到答案。",
    ])
    add_tip(doc, "本手册配套的一键脚本都在 build\\env\\ 目录下（装依赖、建库、启动），新手优先用脚本，熟练后再用命令。")

    # ============ 第1章 环境准备 ============
    add_h(doc, "第 1 章  环境准备", 1)
    add_h(doc, "1.1 电脑配置要求", 2)
    add_table(doc, ["项目", "要求"], [
        ["操作系统", "Windows 10 / 11（64 位）"],
        ["内存", "4 GB 以上（建议 8 GB）"],
        ["可用磁盘", "2 GB 以上"],
        ["网络", "需能下载安装包与 Python 依赖包"],
        ["权限", "本机管理员权限（安装软件用）"],
    ], widths=[4, 11])

    add_h(doc, "1.2 需要安装的软件清单", 2)
    add_table(doc, ["软件", "版本", "用途", "是否必装"], [
        ["Python", "3.8+（实测 3.9.7）", "运行后端服务", "必装"],
        ["MySQL Server", "5.7（实测 5.7.32）", "存储业务数据", "必装"],
        ["Git", "最新版", "拉取/管理代码", "必装"],
        ["Microsoft Word", "2016+", "打开生成的文档", "必装"],
        ["VS Code / PyCharm", "最新版", "查看修改代码", "选装"],
        ["MySQL Workbench", "最新版", "图形化看数据库", "选装"],
    ], widths=[3.5, 3.5, 5, 3])
    add_tip(doc, "各软件的官方下载地址与安装要点，见 build\\env\\工具清单.md。")

    add_h(doc, "1.3 本手册与代码的位置", 2)
    add_p(doc, "本手册位于项目 build 目录下，配套脚本位于 build/env 目录：")
    add_code(doc,
             "build/\n"
             "  ├─ 5000B管理系统环境搭建手册.docx   ← 你正在看的这份\n"
             "  ├─ README.md                       ← build 目录说明\n"
             "  └─ env/\n"
             "       ├─ check_env.py               ← 环境自检工具\n"
             "       ├─ 01_install_deps.bat        ← 一键装依赖\n"
             "       ├─ 02_init_db.bat             ← 一键建库+导数据\n"
             "       ├─ 03_start_all.bat           ← 一键启动系统\n"
             "       └─ 工具清单.md                 ← 软件下载地址")

    # ============ 第2章 安装 Python ============
    add_h(doc, "第 2 章  安装 Python", 1)
    add_steps(doc, [
        "打开官方下载页：https://www.python.org/downloads/release/python-3913/",
        "页面拉到最底部，选择 Windows installer (64-bit) 下载。",
        "双击安装包，【务必勾选】最下方的 Add Python to PATH，然后点击 Install Now。",
        "等待安装完成，点击 Close。",
    ])
    add_h(doc, "2.1 验证是否安装成功", 2)
    add_p(doc, "按 Win+R，输入 cmd 回车，在黑窗口里输入：")
    add_code(doc, "python --version")
    add_p(doc, "看到类似 Python 3.9.7 的输出即为成功。")
    add_tip(doc, "若提示“'python' 不是内部或外部命令”，说明安装时没勾选 Add to PATH。解决办法：重新运行安装包 → 选 Modify → 勾选 Add Python to environment variables，或者卸载重装并记得勾选。")

    # ============ 第3章 安装 MySQL ============
    add_h(doc, "第 3 章  安装 MySQL 5.7", 1)
    add_steps(doc, [
        "打开：https://dev.mysql.com/downloads/mysql/5.7.html",
        "选择 Windows (x86, 64-bit) 的 MSI Installer 下载（可跳过登录，点 No thanks, just start my download）。",
        "安装类型选 Server only（只装数据库服务，够用且快）。",
        "一路 Next，端口保持 3306 不要改。",
        "设置 root 用户密码。为了和项目默认配置一致，建议先设成 root。",
        "继续 Next 直到安装完成。",
    ])
    add_h(doc, "3.1 启动 MySQL 服务", 2)
    add_p(doc, "装完通常会自动启动。若不确定，在命令行输入（服务名可能是 MySQL5 / MySQL57，以你机器为准）：")
    add_code(doc, "net start MySQL5")
    add_h(doc, "3.2 验证数据库可连接", 2)
    add_code(doc,
             'mysql --version\n'
             '"C:\\Program Files\\MySQL\\MySQL Server 5.7\\bin\\mysql.exe" -uroot -proot -e "SELECT 1;"')
    add_p(doc, "能输出版本号、并且第二条能返回一个 1，说明数据库可用。")
    add_tip(doc, "若 net start 提示“服务名无效”，先查真实服务名：命令行输入 sc query type= service state= all | findstr MYSQL，用查到的名字替换 MySQL5。")
    add_tip(doc, "若你安装时设了别的密码，请同步修改两处：backend\\config.py 里的 DATABASE_URL，以及 build\\env\\02_init_db.bat 里的 DBPWD。")

    # ============ 第4章 获取代码 ============
    add_h(doc, "第 4 章  获取代码", 1)
    add_steps(doc, [
        "安装 Git：https://git-scm.com/download/win ，一路 Next 即可。",
        "在你想放代码的盘符下（例如 D:\\5000），右键 → Git Bash Here，执行克隆命令：",
    ])
    add_code(doc, "git clone <仓库地址> 5000BManagePro\ncd 5000BManagePro")
    add_p(doc, "若暂时拿不到仓库地址，也可以直接向同事拷贝整个 5000BManagePro 文件夹，效果一样。")
    add_h(doc, "4.1 目录结构速览", 2)
    add_table(doc, ["目录", "内容"], [
        ["backend/", "后端服务（FastAPI 接口、文档生成引擎）"],
        ["frontend/", "前端页面（HTML/CSS/JS）"],
        ["database/", "数据库导出文件 gjb5000b.sql（结构+数据）"],
        ["build/", "环境搭建专区（本手册 + 一键脚本）"],
        ["scripts/", "数据初始化、导入等维护脚本"],
        ["start.bat", "一键启动后端+前端（最常用）"],
        ["requirements.txt", "后端 Python 依赖清单"],
    ], widths=[4, 11])

    # ============ 第5章 安装后端依赖 ============
    add_h(doc, "第 5 章  安装后端依赖", 1)
    add_p(doc, "最简单的方式：双击运行 build\\env\\01_install_deps.bat，等待出现“依赖安装完成”。")
    add_p(doc, "也可以手动执行（在项目根目录的命令行里）：")
    add_code(doc, "python -m pip install -r requirements.txt -i https://pypi.tuna.tsinghua.edu.cn/simple")
    add_tip(doc, "加 -i 清华镜像是为了下载快。若公司网络屏蔽了镜像，去掉 -i 及后面的网址用官方源，或换手机热点。")
    add_tip(doc, "若报权限错误，在命令后加 --user；若依赖冲突，可用虚拟环境：python -m venv .venv 然后 .venv\\Scripts\\activate。")

    # ============ 第6章 初始化数据库 ============
    add_h(doc, "第 6 章  初始化数据库", 1)
    add_p(doc, "最简单的方式：双击运行 build\\env\\02_init_db.bat（自动建库 + 导入数据 + 列出表）。")
    add_p(doc, "手动执行的等价命令：")
    add_code(doc,
             'set MYSQL="C:\\Program Files\\MySQL\\MySQL Server 5.7\\bin\\mysql.exe"\n'
             '%MYSQL% -uroot -proot -e "CREATE DATABASE IF NOT EXISTS gjb5000b DEFAULT CHARACTER SET utf8mb4 COLLATE utf8mb4_general_ci;"\n'
             '%MYSQL% -uroot -proot --default-character-set=utf8mb4 gjb5000b < database\\gjb5000b.sql\n'
             '%MYSQL% -uroot -proot --default-character-set=utf8mb4 gjb5000b -e "SHOW TABLES;"')
    add_tip(doc, "最后一条 SHOW TABLES 能列出很多表名，就说明数据导入成功了。若提示 Unknown database，一定是前面的建库步骤没执行成功。")

    # ============ 第7章 启动系统 ============
    add_h(doc, "第 7 章  启动系统", 1)
    add_p(doc, "双击项目根目录下的 start.bat，它会自动启动后端与前端，并提示启动结果。")
    add_table(doc, ["服务", "地址", "说明"], [
        ["后端接口", "http://127.0.0.1:8000", "接口文档加 /docs，如 http://127.0.0.1:8000/docs"],
        ["前端页面", "http://127.0.0.1:8080", "浏览器打开这个地址使用系统"],
    ], widths=[3, 6, 6])
    add_p(doc, "手动启动方式（排错时用）：")
    add_code(doc, "python run_backend.py\npython scripts/frontend_server.py")
    add_tip(doc, "务必用 start.bat 一键启停，不要自己另开窗口启动多个后端，否则容易出现端口被旧进程占用、页面一直转圈的问题。")

    # ============ 第8章 登录与界面 ============
    add_h(doc, "第 8 章  登录与界面总览", 1)
    add_steps(doc, [
        "浏览器打开 http://127.0.0.1:8080 ，出现登录页。",
        "登录后进入主界面，左侧为功能菜单。",
        "顶栏可查看/切换当前项目，也可以「＋ 新建项目」「⇄ 切换项目」「修改项目」。",
    ])
    add_table(doc, ["菜单", "作用"], [
        ["PP 项目策划", "项目信息、进度、风险、资源、估算、相关方，以及**生成文档**"],
        ["PMC 项目监控", "项目监控数据"],
        ["告警", "告警信息查看与处理"],
        ["基础数据", "需求、风险、相关方等基础配置"],
        ["用户管理", "账号与角色"],
        ["模板中心", "文档模板查看"],
        ["系统设置", "数据源、SVN 等配置"],
    ], widths=[4, 11])

    # ============ 第9章 生成第一份开发计划 ============
    add_h(doc, "第 9 章  生成第一份《软件开发计划》", 1)
    add_steps(doc, [
        "进入「PP 项目策划」页，确认顶栏当前项目（例如 R105）。",
        "找到「下载到本地 / 生成」按钮并点击。",
        "在弹窗里选择保存目录（例如 D:\\5000\\R105），填写文件名（例如 R105_SDP.docx）。",
        "点击确定，等待进度条走完，出现绿色「下载完成」。",
        "到刚才选择的目录，用 Word 打开生成的 docx 文件。",
    ])
    add_h(doc, "9.1 生成文档里的只读区域（重要）", 2)
    add_p(doc, "平台自动生成的数据表（进度、风险、相关方、软硬件资源、文档/代码规模、数据管理等共 10 张表）在 Word 里是**不可编辑**的，这是刻意保护的平台数据；正文叙述与需要人工填写的表格则可以正常编辑。")
    add_tip(doc, "如果你发现整篇都能编辑，请先确认打开的到底是不是刚生成的那个文件（看文件修改时间与大小）。老版本文件或别人拷来的旧文件是没有这层保护的。")

    add_h(doc, "9.2 生成失败：提示字段未填写", 2)
    add_p(doc, "系统会在生成前校验项目关键字段。若缺少，会直接报错并列出缺哪些，例如：")
    add_code(doc, "生成《软件开发计划》失败：以下关键字段未填写，请先在「项目信息」中补全 —— 顾客代表单位、批准日期、承研单位")
    add_p(doc, "处理办法：点顶栏「修改项目」，把提示的字段补齐后重新生成即可。这是为了防止文档里出现空值或占位默认值。")

    # ============ 第10章 备份与恢复 ============
    add_h(doc, "第 10 章  数据库备份与恢复", 1)
    add_p(doc, "备份（导出当前数据库到 database\\gjb5000b.sql）：")
    add_code(doc,
             '"C:\\Program Files\\MySQL\\MySQL Server 5.7\\bin\\mysqldump.exe" --user=root --password=root '
             '--host=127.0.0.1 --port=3306 --default-character-set=utf8mb4 --single-transaction '
             '--routines --events --result-file=database/gjb5000b.sql gjb5000b')
    add_p(doc, "恢复（把 sql 导回数据库，即第 6 章的导入步骤）：")
    add_code(doc,
             '"C:\\Program Files\\MySQL\\MySQL Server 5.7\\bin\\mysql.exe" -uroot -proot '
             '--default-character-set=utf8mb4 gjb5000b < database\\gjb5000b.sql')
    add_tip(doc, "项目约定：每次提交代码到 Git 前，都要重新导出一次 database\\gjb5000b.sql 并一起提交，保证别人拉到代码能直接还原出完整数据。")

    # ============ 第11章 常见问题排查 ============
    add_h(doc, "第 11 章  常见问题排查", 1)
    add_table(doc, ["现象", "可能原因", "处理办法"], [
        ["页面一直转圈 / 卡处理中", "后端没启动，或旧进程（孤儿 worker）占了端口",
         "关闭所有后端窗口，用 start.bat 重新启动；仍不行则重启电脑后重试"],
        ["提示 Failed to fetch", "后端服务没起来，或端口 8000 被别的程序占用",
         "先访问 http://127.0.0.1:8000/docs 看能否打开；打不开就是后端没起来"],
        ["8000 端口被占用", "上次启动的进程没退干净",
         "netstat -ano | findstr :8000 找到 PID，再 taskkill /PID <PID> /F"],
        ["数据库连不上 / Access denied", "MySQL 服务没启动，或密码不对",
         "net start MySQL5；确认 backend\\config.py 里的账号密码与实际一致"],
        ["Unknown database 'gjb5000b'", "还没建库",
         "运行 build\\env\\02_init_db.bat 建库并导入数据"],
        ["pip 安装很慢或超时", "网络访问官方源慢",
         "命令后加 -i https://pypi.tuna.tsinghua.edu.cn/simple；或换网络"],
        ["'python' 不是内部或外部命令", "安装时没勾选 Add to PATH",
         "重装 Python 并勾选，或手动把 Python 目录加入系统环境变量 PATH"],
        ["生成的 Word 打开是空的", "打开的不是新生成的文件（旧文件/别人拷来的文件）",
         "确认文件修改时间与大小；重新生成后立即打开新文件"],
        ["保存利益相关方报错", "填写内容超过字段长度",
         "联系开发确认字段长度；一般缩短填写内容即可"],
        ["改了代码但没生效", "后端进程没重启（跑的还是旧代码）",
         "必须关闭后端再启动（用 start.bat 一键重启），浏览器 Ctrl+F5 强刷"],
    ], widths=[4.5, 5, 5.5])

    # ============ 第12章 日常维护 ============
    add_h(doc, "第 12 章  日常维护", 1)
    add_bullets(doc, [
        "停止服务：直接关闭启动出来的黑窗口，或在 start.bat 里选择停止。",
        "查看日志：项目 logs\\ 目录，以及项目根的 _start_out.txt（启动输出）。",
        "清理残留进程：start.bat 已内置清理逻辑，务必用它启停，不要手动零散启动。",
        "更新代码：git pull 拉取最新代码后，重启后端才会生效；若数据库有更新，重新导入 gjb5000b.sql。",
        "环境自检：随时可运行 python build\\env\\check_env.py，逐项查看 [OK] / [失败]。",
    ])

    # ============ 附录 ============
    add_h(doc, "附录 A  常用命令速查", 1)
    add_code(doc,
             "python --version                     查看 Python 版本\n"
             "python build\\env\\check_env.py        环境自检\n"
             "net start MySQL5                     启动 MySQL 服务\n"
             "netstat -ano | findstr :8000         查看 8000 端口占用\n"
             "taskkill /PID <PID> /F               强制结束指定进程\n"
             "git pull                             拉取最新代码\n"
             "python -m pip install -r requirements.txt -i https://pypi.tuna.tsinghua.edu.cn/simple")

    add_h(doc, "附录 B  端口与账号速查", 1)
    add_table(doc, ["项目", "值"], [
        ["后端地址", "http://127.0.0.1:8000（接口文档 /docs）"],
        ["前端地址", "http://127.0.0.1:8080"],
        ["数据库地址", "127.0.0.1:3306"],
        ["数据库名", "gjb5000b"],
        ["数据库账号", "root / root（默认，见 backend\\config.py）"],
        ["数据库导出文件", "database\\gjb5000b.sql"],
        ["一键启动", "项目根 start.bat"],
    ], widths=[4.5, 10.5])

    add_h(doc, "附录 C  寻求帮助", 1)
    add_bullets(doc, [
        "先运行 python build\\env\\check_env.py，把 [失败] 项截图。",
        "把报错信息完整复制（不要只截一部分），连同操作步骤一起反馈。",
        "说明你进行到第几章、执行了哪条命令，便于快速定位。",
    ])

    doc.save(OUT)
    print("MANUAL_SAVED", OUT)
    print("SIZE", os.path.getsize(OUT))
    print("PARAGRAPHS", len(doc.paragraphs))
    print("TABLES", len(doc.tables))


if __name__ == "__main__":
    main()
